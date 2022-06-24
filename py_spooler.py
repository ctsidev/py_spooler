import csv
import cx_Oracle as ora
import json, sys, os
from docx import Document
from datetime import date

username = input('Enter ClarityQ schema username:\n')
password = input(f'Password for {username} schema:\n')

doc = Document()
try:
    cur = ora.connect(f'{username}/{password}@clarityq').cursor()
except:
    print('\nCould not connect to ClarityQ, quitting..')
else:
    print('\nConnected to ClarityQ')
    print('--------------------------------')

def header_element(element_name, element_desc):
    doc.add_heading(element_name, level=3)
    doc.add_paragraph(element_desc)

def get_element_info(element):
    sql = """
    SELECT FIELD_NAME,
    DESCRIPTION,
    COMMENT_FIELD 
    FROM ctsi_research.JS_DATA_DICTIONARY
    WHERE ELEMENT =: element
    """
    cur.execute(sql,{'element':str(element)})
    output = {}
        
    for item in cur:
        output[str(item[0]).upper()] = [str(item[1]),str(item[2])]

    return output  
    
def create_table(data_element_entries):
    table = doc.add_table(rows=1, cols=3, style = 'Light Grid Accent 4')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field Name'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Comment'

    for item in data_element_entries:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item[0])
        row_cells[1].text = str(item[1])
        row_cells[2].text = str(item[2]) 

def main():
    project_id = input('Enter project_id:\n')
    folder = input('Enter project directory:\n')
    selection = int(input('1) Spool\n2) Data Dictionary\n3) Both\n\nEnter Selection:\n'))

    if selection == 1 or selection == 3:
        try:
            f = open(f'{folder}/xdr_{project_id}_SPOOL.sql')
        except:
            print('Directory not found')
            exit()
            
        statements = f.read().split(';')

        print('--------------------------------')

        for statement in statements:
            if statement.strip():
                statement = statement.strip()
                filename = statement.split('\n')[0][3:]
                file_loc = f"{folder}/Data/{filename}"
                print(f'Starting spool: {file_loc}')
                print(f'Executing query..')

                try:
                    db_data = cur.execute(statement)
                    cols = [row[0] for row in db_data.description]

                    print('Data ready, creating file..')
                    csv_file = open(file_loc, "w")

                    delimiter = ','
                    if filename == 'HIPAA.txt':
                        delimiter = ';'
                        
                    writer = csv.writer(csv_file, delimiter=delimiter, lineterminator="\n", quoting=csv.QUOTE_NONNUMERIC)
                    writer.writerow(cols)
                    for row in db_data:
                        writer.writerow(row)

                    csv_file.close()
                    print('File created')
                    print('--------------------------------')
                except ora.DatabaseError as exc:
                    error, = exc.args
                    print(f'Query failed, {error.code}: {error.message}, skipping')
                    print('--------------------------------')

    if selection == 2 or selection == 3:
        print('Creating Data Dictionary..')
        sql = """
        SELECT INVESTIGATOR
            ,IRB
            ,DESCRIPTION
            ,to_char(LATEST_FROM_DATE, 'mm/dd/yyyy') as LATEST_FROM_DATE
            ,to_char(LATEST_TO_DATE, 'mm/dd/yyyy') as LATEST_TO_DATE
        FROM I2B2.BIP_PROJECT
        WHERE project_id =: project_id
        """

        cur.execute(sql, {'project_id': project_id})
        for item in cur:
            irb = str(item[1])
            investigator = str(item[0])
            description = str(item[2])
            from_date = str(item[3])
            to_date = str(item[4])

        ## CREATE DOCUMENT OBJECT ##

        doc.add_heading('Data Dictionary', level = 1)
        doc.add_heading('PI: ' + investigator, level = 2)
        doc.add_heading('IRB#: '  + irb, level = 2)
        header_element('Background: ', description)  
        header_element('Extraction Date: ', str(date.today().strftime("%m/%d/%Y")))
        header_element('Extraction Timeframe: ', f'{from_date} - {to_date}')
        header_element('Selection Criteria: ', '')

        ## GET ELEMENTS ##

        sql = """
        SELECT element, description
        FROM ctsi_research.JS_DATA_DICTIONARY_MAIN 
        ORDER BY order_num
        """
        cur.execute(sql)
        dd_elements = [el for el in cur]
        element_names = [el[0] for el in dd_elements]
        data_files = [f for f in os.listdir(f'{folder}\Data') if f.endswith('.csv')]

        print('Checking CSV files..')

        # Delete first line of each CSV if first line blank
        for filename in data_files:
            if filename[:-4] not in element_names:
                filename = filename[:-4]
                dd_elements.append((filename, '[Ad hoc entity description]'))

        print('Preparing Data Dictionary..')
        for element in dd_elements:
            f = f'{folder}\Data\{element[0]}.csv'
            if os.path.exists(f):
                doc.add_heading(element[0], level=3)
                doc.add_paragraph(element[1])

                headers = open(f).readline().replace('\n','').replace('"','').split(',')
                element_info = get_element_info(element[0])
                data_element_entries = []

                for header in headers:
                    header = str(header).upper()
                    if header in element_info:
                        rows = (header, element_info[header][0], element_info[header][1])
                    else:
                        if header == 'IP_PATIENT_ID':
                            rows = (header, 'A de-identified, unique ID number for the individual patient (assigned by IP for coding purposes)', 'Use this variable to link to variables from all other files that contain patient information')
                        elif header == 'IP_ENC_ID':
                            rows = (header, 'A de-identified, unique ID number for the encounter (assigned by IP for coding purposes)', 'Use this variable to link to variables from other tables with encounter information e.g. diagnoses, procedures, vital signs, etc.')
                        else:
                            rows = (header, '[Description for ad hoc field]' , '[Comment for ad hoc field]')
                            
                    data_element_entries.append(rows)

                create_table(data_element_entries)

        ## WRITE FILE ##

        doc.save(f'{folder}/DataDictionary_{investigator}_#{irb}_.docx')

        print('Data dictionary successfully created at:')
        print(f'{folder}/DataDictionary_{investigator}_#{irb}_.docx')

if __name__ == '__main__':
    main()